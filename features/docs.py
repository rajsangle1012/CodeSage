import streamlit as st
import os
import io
from groq import Groq
from utils import get_supabase, load_css

# --- EXPORT LIBRARIES ---
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def create_pdf(text_content):
    """Generates a PDF file from the text content."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=11)
    
    clean_text = text_content.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "CodeSage Technical Documentation", ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 7, clean_text)
    
    return pdf.output(dest='S').encode('latin-1')

def create_docx(text_content):
    """Generates a Microsoft Word (.docx) file."""
    doc = Document()
    doc.add_heading('CodeSage Technical Documentation', 0)
    doc.add_paragraph(text_content)
    
    io_stream = io.BytesIO()
    doc.save(io_stream)
    return io_stream.getvalue()

def generate_docs(input_text, language="Python"):
    """Uses Groq AI to generate both high-level docs and inline comments."""
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return "⚠️ Error: GROQ_API_KEY not found."
        
    client = Groq(api_key=api_key)
    
    system_prompt = f"""
    You are an elite Software Architect and Technical Writer. 
    The user will provide a {language} code snippet.
    
    You MUST provide TWO things in your response, formatted in Markdown:
    
    1. ### 📖 Technical Documentation
       Provide a highly professional summary of what the code does. Include:
       - **Purpose:** Brief description.
       - **Parameters/Inputs:** What it expects.
       - **Returns/Outputs:** What it yields.
       - **Usage Example:** A quick example of how to call it.
       
    2. ### 💻 Documented Code
       Provide the COMPLETE original code, but heavily upgraded with highly professional, explanatory inline comments and standard docstrings. 
       Do not alter the core logic, ONLY add documentation.
       Wrap the code in standard Markdown code blocks (```).
    """
        
    try:
        with st.spinner("CodeSage is writing your documentation... ✍️"):
            completion = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": system_prompt},
                          {"role": "user", "content": f"DOCUMENT THIS CODE:\n{input_text}"}],
                temperature=0.2
            )
            return completion.choices[0].message.content or ""
    except Exception as e:
        return f"❌ System Error: {str(e)}"

def render_docs_page():
    load_css()
    
    # --- CSS STYLING ---
    st.markdown("""
        <style>
            .cyber-header {
                font-family: 'JetBrains Mono', monospace; 
                font-size: 3.2rem; 
                font-weight: 800; 
                background: linear-gradient(135deg, #8B5CF6, #fff);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                text-shadow: 0 0 20px rgba(139, 92, 246, 0.5);
                text-align: center;
                text-transform: uppercase;
                margin-bottom: 0px;
                animation: textGlowPurple 3s infinite alternate;
            }
            @keyframes textGlowPurple {
                0% { text-shadow: 0 0 10px rgba(139, 92, 246, 0.4); }
                100% { text-shadow: 0 0 25px rgba(139, 92, 246, 0.8), 0 0 40px rgba(139, 92, 246, 0.4); }
            }
            .cyber-sub {
                text-align: center; color: #94A3B8; letter-spacing: 3px; font-family: monospace; font-size: 1rem; margin-bottom: 30px;
            }

            .tool-icon {
                display: inline-block;
                animation: toolActivity 2s infinite ease-in-out;
                -webkit-background-clip: border-box !important;
                -webkit-text-fill-color: initial !important;
            }
            @keyframes toolActivity {
                0%, 100% { transform: scale(1) rotate(0deg); filter: drop-shadow(0 0 10px rgba(139, 92, 246, 0.4)); }
                50% { transform: scale(1.15) rotate(-5deg); filter: drop-shadow(0 0 25px rgba(139, 92, 246, 1)); }
            }

            .glow-wrapper {
                border: 1px solid rgba(139, 92, 246, 0.4);
                box-shadow: 0 0 20px rgba(139, 92, 246, 0.3);
                padding: 30px;
                border-radius: 12px;
                background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, rgba(59, 130, 246, 0.1) 100%);
                margin-bottom: 30px;
                position: relative;
                overflow: hidden;
            }
            .glow-wrapper::after {
                content: ""; position: absolute; top: 0; left: 0; width: 100%; height: 100%;
                background: linear-gradient(to right, transparent, rgba(139, 92, 246, 0.15), transparent);
                transform: translateX(-100%);
                animation: scanRightPurple 3s linear infinite;
                pointer-events: none;
            }
            @keyframes scanRightPurple { 100% { transform: translateX(100%); } }

            .editor-label { 
                font-size: 0.9rem; 
                color: #8B5CF6; 
                font-family: 'JetBrains Mono', monospace; 
                font-weight: 800; 
                text-transform: uppercase; 
                margin-bottom: 5px; 
                letter-spacing: 2px; 
                animation: pulseTextPurple 3s infinite alternate; 
            }
            @keyframes pulseTextPurple { 0% { opacity: 0.8; text-shadow: 0 0 5px rgba(139, 92, 246, 0.3); } 100% { opacity: 1; text-shadow: 0 0 15px rgba(139, 92, 246, 0.8); } }

            div[data-baseweb="select"] > div {
                background: rgba(15, 23, 42, 0.8) !important;
                border: 2px solid rgba(139, 92, 246, 0.4) !important;
                border-radius: 8px !important;
                box-shadow: 0 0 15px rgba(139, 92, 246, 0.1);
                color: #A78BFA !important;
            }

            .stTextArea textarea { 
                background: linear-gradient(180deg, rgba(2,6,23,0.9), rgba(15,23,42,0.8)) !important;
                border: 1px solid rgba(139, 92, 246, 0.3) !important; 
                border-left: 4px solid #3B82F6 !important;
                color: #E2E8F0 !important; 
                font-family: 'JetBrains Mono', monospace !important; 
                border-radius: 8px !important; 
                box-shadow: inset 0 0 20px rgba(0,0,0,0.8);
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            }
            .stTextArea textarea:focus { 
                border-color: #8B5CF6 !important; 
                border-left: 4px solid #8B5CF6 !important;
                box-shadow: 0 0 25px rgba(139, 92, 246, 0.3), inset 0 0 15px rgba(139, 92, 246, 0.1) !important; 
                transform: scale(1.01);
            }
            
            div[data-testid="stVerticalBlockBorderWrapper"] {
                background: linear-gradient(180deg, rgba(2,6,23,0.9), rgba(15,23,42,0.8));
                border: 1px solid rgba(139, 92, 246, 0.3);
                border-left: 4px solid #3B82F6;
                border-radius: 8px;
                box-shadow: inset 0 0 20px rgba(0,0,0,0.8);
                color: #E2E8F0;
                overflow-y: auto;
            }

            .export-tools {
                background: rgba(15, 23, 42, 0.6);
                padding: 15px;
                border-radius: 8px;
                border-left: 4px solid #3B82F6;
                box-shadow: 0 5px 15px rgba(0,0,0,0.5);
                margin-top: 15px;
            }

            [data-testid="stFileUploader"] > div {
                background: rgba(2, 6, 23, 0.6) !important;
                border: 1px dashed rgba(139, 92, 246, 0.5) !important;
                border-radius: 8px;
                transition: all 0.3s;
            }
            [data-testid="stFileUploader"] > div:hover {
                border-color: #3B82F6 !important;
                background: rgba(59, 130, 246, 0.05) !important;
                box-shadow: 0 0 15px rgba(59, 130, 246, 0.3);
            }

            button[kind="primary"] {
                background: linear-gradient(90deg, #8B5CF6, #3B82F6, #8B5CF6) !important;
                background-size: 200% auto !important;
                color: #fff !important;
                font-weight: 900 !important;
                border: none !important;
                animation: gradientShiftPurple 3s linear infinite !important;
                transition: all 0.3s !important;
                text-transform: uppercase;
                letter-spacing: 2px;
                box-shadow: 0 0 20px rgba(139, 92, 246, 0.4) !important;
            }
            button[kind="primary"]:hover {
                transform: scale(1.01) !important;
                box-shadow: 0 0 30px rgba(139, 92, 246, 0.6) !important;
            }
            @keyframes gradientShiftPurple { 0% { background-position: 0% center; } 100% { background-position: 200% center; } }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("<h1 class='cyber-header'><span class='tool-icon'>📝</span> Auto-Docs Engine</h1>", unsafe_allow_html=True)
    st.markdown("<p class='cyber-sub'>UPLOAD A SCRIPT TO AUTOGENERATE INLINE COMMENTS AND TECHNICAL SPECS.</p>", unsafe_allow_html=True)
    st.divider()

    if not HAS_FPDF or not HAS_DOCX:
        st.warning("⚠️ For PDF and Word exports, run `pip install fpdf python-docx` in your terminal.")

    # --- INPUT SETTINGS ---
    st.markdown('<div class="glow-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="editor-label">🎯 Target Configuration</div>', unsafe_allow_html=True)
    c_conf1, c_conf2 = st.columns(2)
    with c_conf1:
        language = st.selectbox("💻 Select Target Language", ["Python", "JavaScript", "Java", "C++", "HTML/CSS", "TypeScript", "SQL"], label_visibility="collapsed")
    with c_conf2:
        uploaded_file = st.file_uploader("Upload a file instead of pasting", type=["py", "js", "java", "cpp", "ts", "html", "css", "txt", "sql"], label_visibility="collapsed")
        
        if uploaded_file is not None:
            if st.session_state.get("last_uploaded_file") != uploaded_file.name:
                try:
                    file_contents = uploaded_file.getvalue().decode("utf-8")
                    st.session_state["doc_input_text"] = file_contents
                    st.session_state["last_uploaded_file"] = uploaded_file.name
                    st.toast(f"File '{uploaded_file.name}' loaded successfully!", icon="✅")
                except Exception as e:
                    st.error(f"Could not read file. Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.write("<br>", unsafe_allow_html=True)

    # --- TEXT WORKSPACE ---
    c_input, c_output = st.columns(2)
    
    with c_input:
        st.markdown('<div class="editor-label">🔴 Undocumented Source</div>', unsafe_allow_html=True)
        default_text = st.session_state.get("doc_input_text", "")
        input_text = st.text_area("input", value=default_text, height=550, label_visibility="collapsed", placeholder="Paste your undocumented code here, or upload a file above...")
        st.session_state["doc_input_text"] = input_text

    with c_output:
        st.markdown('<div class="editor-label">🟢 Generated Technical Specs</div>', unsafe_allow_html=True)
        output_content = st.session_state.get("generated_docs", "")
        
        if output_content:
            with st.container(border=True, height=550):
                st.markdown(output_content)
        else:
            st.text_area("output", value="Your documented code and technical specs will appear here...", height=550, disabled=True, label_visibility="collapsed")

    st.write("<br>", unsafe_allow_html=True)

    # Execution Action Flow
    if st.button("🪄 GENERATE FULL DOCUMENTATION", type="primary", use_container_width=True):
        if not input_text.strip():
            st.warning("⚠️ Please provide code or upload a file first!")
        else:
            generated_content = generate_docs(input_text, language)
            st.session_state["generated_docs"] = generated_content
            
            supabase = get_supabase()
            uid = st.session_state.get("user_id")
            if supabase and uid:
                try:
                    p_res = supabase.table("profiles").select("lifetime_reviews").eq("id", uid).single().execute()
                    if p_res.data:
                        new_xp = p_res.data.get("lifetime_reviews", 0) + 2
                        supabase.table("profiles").update({"lifetime_reviews": new_xp}).eq("id", uid).execute()
                        prof_refresh = supabase.table("profiles").select("*").eq("id", uid).single().execute()
                        st.session_state["profile"] = prof_refresh.data
                except Exception as e: pass
            
            st.rerun()

    if output_content:
        st.markdown('<div class="export-tools">', unsafe_allow_html=True)
        st.markdown("**📥 EXPORT SYSTEM:**")
        dl1, dl2, dl3 = st.columns(3)
        with dl1:
            st.download_button(label="📄 .MD", data=output_content, file_name="CodeSage_Docs.md", mime="text/markdown", use_container_width=True)
        with dl2:
            if HAS_DOCX:
                docx_bytes = create_docx(output_content)
                st.download_button(label="📘 .DOCX", data=docx_bytes, file_name="CodeSage_Docs.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            else:
                st.button("📘 .DOCX", disabled=True, help="Run pip install python-docx")
        with dl3:
            if HAS_FPDF:
                pdf_bytes = create_pdf(output_content)
                st.download_button(label="📕 .PDF", data=pdf_bytes, file_name="CodeSage_Docs.pdf", mime="application/pdf", use_container_width=True)
            else:
                st.button("📕 .PDF", disabled=True, help="Run pip install fpdf")
        st.markdown('</div>', unsafe_allow_html=True)

    # --- NAVIGATION ---
    st.write("<br><br>", unsafe_allow_html=True)
    if st.button("← Return to Dashboard", use_container_width=True):
        st.session_state["page"] = "home"
        st.rerun()